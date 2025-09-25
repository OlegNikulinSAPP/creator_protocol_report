import wx
import os
import sys
import re
from datetime import datetime

# Убедимся, что импортируем правильную библиотеку python-docx
try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    print("Ошибка: Не удалось импортировать python-docx. Установите: pip install python-docx")
    sys.exit(1)


# Функции для работы с протоколами
def extract_protocol_data(filename):
    """Извлекает номер протокола и дату из документа."""
    try:
        doc = Document(filename)
        for paragraph in doc.paragraphs:
            match = re.match(r'ПРОТОКОЛ №(\d+) от (\d{2}\.\d{2}\.\d{4})г?', paragraph.text)
            if match:
                return match.groups()
        return None, None
    except Exception as e:
        print(f"Ошибка при извлечении данных протокола: {e}")
        return None, None


def check_deadline(deadline_str):
    """Проверяет, истёк ли срок мероприятия."""
    today = datetime.now().date()
    try:
        deadline_date = datetime.strptime(deadline_str, "%d.%m.%Y").date()
        return deadline_date <= today
    except ValueError:
        return False


def parse_protocol(filename):
    """Парсит протокол, выбирая первую таблицу для извлечения мероприятий"""
    try:
        doc = Document(filename)

        if not doc.tables:
            print("В документе нет таблиц")
            return []

        first_table = doc.tables[0]
        events = []

        for row in first_table.rows:
            if len(row.cells) > 0:
                cell_text = row.cells[0].text.strip()
                if cell_text:
                    normalized_text = re.sub(r'\s+', ' ', cell_text)
                    events.append(normalized_text)

        return events
    except Exception as e:
        print(f"Ошибка при чтении протокола: {e}")
        return []


def replace_placeholders_in_template(template_doc, protocol_num, date):
    """Замена полей-шаблонов в документе."""
    for paragraph in template_doc.paragraphs:
        paragraph.text = paragraph.text.replace("№…", f"№{protocol_num}").replace("от …г.", f"от {date}г.")


def set_red_font(cell):
    """Устанавливает красный цвет шрифта для текста в ячейке."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет


def create_report(template_file, protocol_file, output_file):
    """Создает отчет на основе шаблона и протокола с табличной структурой"""
    try:
        # Чтение протокола
        events = parse_protocol(protocol_file)
        if not events:
            print("Не удалось извлечь мероприятия из протокола.")
            return False

        # Извлечение номера протокола и даты
        protocol_num, date = extract_protocol_data(protocol_file)
        if not protocol_num or not date:
            print("Не удалось извлечь номер протокола и дату.")
            return False

        # Открытие шаблона
        template_doc = Document(template_file)

        # Заменяем шаблонные маркеры
        replace_placeholders_in_template(template_doc, protocol_num, date)

        # Проверка наличия таблицы в шаблоне
        if not template_doc.tables:
            print("В шаблоне нет таблицы.")
            return False

        # Берем первую таблицу
        table = template_doc.tables[0]

        # Добавление мероприятий в таблицу
        for idx, event in enumerate(events, start=1):
            # Добавляем новую строку
            if idx > 1:  # Первая строка уже существует (заголовки)
                row = table.add_row()
            else:
                # Для первой строки используем существующую (после заголовков)
                if len(table.rows) > 1:
                    row = table.rows[1]
                else:
                    row = table.add_row()

            # Убедимся, что в строке достаточно ячеек
            while len(row.cells) < 2:
                row.add_cell()

            # Заполняем ячейки
            cell = row.cells[0]
            cell.text = f"{idx}. {event}"

            # Проверяем сроки
            deadline_match = re.search(r'Срок:\s*(\d{2}\.\d{2}\.\d{4})', event)
            if deadline_match:
                deadline_str = deadline_match.group(1)
                if check_deadline(deadline_str):
                    set_red_font(cell)
                    row.cells[1].text = "Не выполнено"
                else:
                    row.cells[1].text = "В процессе"
            else:
                row.cells[1].text = ""

        # Сохранение результата
        template_doc.save(output_file)
        print(f"Отчёт успешно сохранён: {output_file}")
        return True

    except Exception as e:
        print(f"Произошла ошибка при создании отчета: {e}")
        return False


class TextRedirector:
    """Перенаправление вывода в текстовый контрол"""

    def __init__(self, text_ctrl):
        self.text_ctrl = text_ctrl

    def write(self, string):
        if string.strip():
            wx.CallAfter(self.text_ctrl.AppendText, string)

    def flush(self):
        pass


class ProtocolApp(wx.Frame):
    def __init__(self):
        super().__init__(None, title="Генератор отчетов по протоколам", size=(800, 600))
        self.protocol_file = ""
        self.output_file = ""
        self.template_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Шаблон.docx")

        self.create_ui()
        self.Centre()
        self.Show()

    def create_ui(self):
        panel = wx.Panel(self)
        main_sizer = wx.BoxSizer(wx.VERTICAL)

        # Заголовок
        title = wx.StaticText(panel, label="Генератор отчетов по протоколам")
        title_font = wx.Font(14, wx.DEFAULT, wx.NORMAL, wx.BOLD)
        title.SetFont(title_font)
        main_sizer.Add(title, 0, wx.ALL | wx.ALIGN_CENTER, 10)

        # Разделитель
        main_sizer.Add(wx.StaticLine(panel), 0, wx.EXPAND | wx.ALL, 5)

        # Выбор файла протокола
        protocol_sizer = wx.BoxSizer(wx.HORIZONTAL)
        protocol_label = wx.StaticText(panel, label="Файл протокола:")
        protocol_sizer.Add(protocol_label, 0, wx.ALL | wx.ALIGN_CENTER_VERTICAL, 5)

        self.protocol_text = wx.TextCtrl(panel, style=wx.TE_READONLY)
        protocol_sizer.Add(self.protocol_text, 1, wx.ALL | wx.EXPAND, 5)

        protocol_btn = wx.Button(panel, label="Выбрать...")
        protocol_btn.Bind(wx.EVT_BUTTON, self.on_select_protocol)
        protocol_sizer.Add(protocol_btn, 0, wx.ALL, 5)
        main_sizer.Add(protocol_sizer, 0, wx.EXPAND | wx.ALL, 5)

        # Информация о шаблоне
        template_sizer = wx.BoxSizer(wx.HORIZONTAL)
        template_label = wx.StaticText(panel, label="Используемый шаблон:")
        template_sizer.Add(template_label, 0, wx.ALL | wx.ALIGN_CENTER_VERTICAL, 5)

        self.template_text = wx.TextCtrl(panel, style=wx.TE_READONLY)
        self.template_text.SetValue(self.template_file)
        template_sizer.Add(self.template_text, 1, wx.ALL | wx.EXPAND, 5)
        main_sizer.Add(template_sizer, 0, wx.EXPAND | wx.ALL, 5)

        # Выбор места сохранения отчета
        output_sizer = wx.BoxSizer(wx.HORIZONTAL)
        output_label = wx.StaticText(panel, label="Сохранить отчет как:")
        output_sizer.Add(output_label, 0, wx.ALL | wx.ALIGN_CENTER_VERTICAL, 5)

        self.output_text = wx.TextCtrl(panel)
        output_sizer.Add(self.output_text, 1, wx.ALL | wx.EXPAND, 5)

        output_btn = wx.Button(panel, label="Выбрать...")
        output_btn.Bind(wx.EVT_BUTTON, self.on_select_output)
        output_sizer.Add(output_btn, 0, wx.ALL, 5)
        main_sizer.Add(output_sizer, 0, wx.EXPAND | wx.ALL, 5)

        # Информация о протоколе
        info_sizer = wx.StaticBoxSizer(wx.VERTICAL, panel, "Информация о протоколе")
        self.info_text = wx.StaticText(panel, label="Выберите файл протокола для отображения информации")
        info_sizer.Add(self.info_text, 0, wx.ALL | wx.EXPAND, 10)
        main_sizer.Add(info_sizer, 0, wx.EXPAND | wx.ALL, 5)

        # Кнопки действий
        button_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.preview_btn = wx.Button(panel, label="Предпросмотр мероприятий")
        self.preview_btn.Bind(wx.EVT_BUTTON, self.on_preview)
        self.preview_btn.Disable()
        button_sizer.Add(self.preview_btn, 0, wx.ALL, 5)

        self.generate_btn = wx.Button(panel, label="Сгенерировать отчет")
        self.generate_btn.Bind(wx.EVT_BUTTON, self.on_generate)
        self.generate_btn.Disable()
        button_sizer.Add(self.generate_btn, 0, wx.ALL, 5)
        main_sizer.Add(button_sizer, 0, wx.ALIGN_CENTER | wx.ALL, 10)

        # Лог действий
        log_sizer = wx.StaticBoxSizer(wx.VERTICAL, panel, "Лог действий")
        self.log_text = wx.TextCtrl(panel, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_RICH2)
        log_sizer.Add(self.log_text, 1, wx.ALL | wx.EXPAND, 5)
        main_sizer.Add(log_sizer, 1, wx.EXPAND | wx.ALL, 5)

        panel.SetSizer(main_sizer)
        self.check_template_exists()

    def check_template_exists(self):
        """Проверяет наличие файла шаблона"""
        if not os.path.exists(self.template_file):
            wx.MessageBox(
                f"Файл шаблона не найден:\n{self.template_file}\n\n"
                "Пожалуйста, убедитесь, что файл 'Шаблон.docx' находится в корневой папке приложения.",
                "Файл шаблона не найден", wx.OK | wx.ICON_WARNING
            )
            self.log_message(f"ВНИМАНИЕ: Файл шаблона не найден: {self.template_file}")
        else:
            self.log_message(f"Шаблон загружен: {self.template_file}")

    def on_select_protocol(self, event):
        with wx.FileDialog(self, "Выберите файл протокола",
                           wildcard="Word files (*.docx)|*.docx",
                           style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST) as dialog:
            if dialog.ShowModal() == wx.ID_CANCEL:
                return
            self.protocol_file = dialog.GetPath()
            self.protocol_text.SetValue(self.protocol_file)
            self.update_protocol_info()
            self.check_buttons_state()

    def on_select_output(self, event):
        default_name = "Отчет_о_исполнении_протокола.docx"
        if self.protocol_file:
            protocol_name = os.path.splitext(os.path.basename(self.protocol_file))[0]
            default_name = f"Отчет_{protocol_name}.docx"

        with wx.FileDialog(self, "Сохранить отчет как", defaultFile=default_name,
                           wildcard="Word files (*.docx)|*.docx",
                           style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dialog:
            if dialog.ShowModal() == wx.ID_CANCEL:
                return
            self.output_file = dialog.GetPath()
            if not self.output_file.endswith('.docx'):
                self.output_file += '.docx'
            self.output_text.SetValue(self.output_file)
            self.check_buttons_state()

    def update_protocol_info(self):
        if not self.protocol_file:
            self.info_text.SetLabel("Выберите файл протокола для отображения информации")
            return

        try:
            protocol_num, date = extract_protocol_data(self.protocol_file)
            events = parse_protocol(self.protocol_file)

            if protocol_num and date:
                info = f"Протокол №{protocol_num} от {date}\nНайдено мероприятий: {len(events)}"

                # Проверяем просроченные мероприятия
                overdue_count = 0
                for event in events:
                    deadline_match = re.search(r'Срок:\s*(\d{2}\.\d{2}\.\d{4})', event)
                    if deadline_match:
                        deadline_str = deadline_match.group(1)
                        if check_deadline(deadline_str):
                            overdue_count += 1

                if overdue_count > 0:
                    info += f" (просрочено: {overdue_count})"

                self.info_text.SetLabel(info)
                self.log_message(f"Загружен протокол №{protocol_num} от {date}, мероприятий: {len(events)}")
            else:
                self.info_text.SetLabel("Не удалось извлечь данные из протокола")
                self.log_message("Ошибка: не удалось извлечь данные из протокола")

        except Exception as e:
            self.info_text.SetLabel(f"Ошибка при чтении протокола: {str(e)}")
            self.log_message(f"Ошибка при чтении протокола: {str(e)}")

    def on_preview(self, event):
        if not self.protocol_file:
            return
        try:
            events = parse_protocol(self.protocol_file)
            self.show_events_preview(events)
        except Exception as e:
            wx.MessageBox(f"Ошибка при предпросмотре мероприятий: {str(e)}", "Ошибка", wx.OK | wx.ICON_ERROR)

    def show_events_preview(self, events):
        dialog = wx.Dialog(self, title="Предпросмотр мероприятий", size=(600, 400))
        panel = wx.Panel(dialog)
        sizer = wx.BoxSizer(wx.VERTICAL)

        label = wx.StaticText(panel, label=f"Найдено мероприятий: {len(events)}")
        sizer.Add(label, 0, wx.ALL, 10)

        listbox = wx.ListBox(panel, choices=events, style=wx.LB_SINGLE)
        sizer.Add(listbox, 1, wx.ALL | wx.EXPAND, 10)

        ok_btn = wx.Button(panel, label="OK")
        ok_btn.Bind(wx.EVT_BUTTON, lambda e: dialog.EndModal(wx.ID_OK))
        sizer.Add(ok_btn, 0, wx.ALIGN_CENTER | wx.ALL, 10)

        panel.SetSizer(sizer)
        dialog.ShowModal()

    def on_generate(self, event):
        if not all([self.protocol_file, self.output_file]):
            wx.MessageBox("Выберите файл протокола и укажите место сохранения отчета!", "Внимание",
                          wx.OK | wx.ICON_WARNING)
            return

        if not os.path.exists(self.template_file):
            wx.MessageBox(f"Файл шаблона не найден!\n\n{self.template_file}", "Ошибка", wx.OK | wx.ICON_ERROR)
            return

        try:
            original_stdout = sys.stdout
            sys.stdout = TextRedirector(self.log_text)

            self.log_message("Начало генерации отчета...")
            self.generate_btn.Disable()

            success = create_report(self.template_file, self.protocol_file, self.output_file)

            if success:
                self.log_message("Генерация отчета завершена успешно!")
                result_dialog = wx.MessageDialog(self, f"Отчет успешно сохранен!\n\nХотите открыть файл?", "Успех",
                                                 wx.YES_NO | wx.ICON_INFORMATION)
                if result_dialog.ShowModal() == wx.ID_YES:
                    os.startfile(self.output_file)
            else:
                self.log_message("Ошибка при генерации отчета!")

        except Exception as e:
            error_msg = f"Ошибка при генерации отчета: {str(e)}"
            self.log_message(error_msg)
            wx.MessageBox(error_msg, "Ошибка", wx.OK | wx.ICON_ERROR)
        finally:
            sys.stdout = original_stdout
            self.generate_btn.Enable()

    def check_buttons_state(self):
        has_protocol = bool(self.protocol_file)
        has_output = bool(self.output_file)
        has_template = os.path.exists(self.template_file)

        self.preview_btn.Enable(has_protocol)
        self.generate_btn.Enable(has_protocol and has_output and has_template)

    def log_message(self, message):
        timestamp = wx.DateTime.Now().FormatTime() + " " + wx.DateTime.Now().FormatDate()
        self.log_text.AppendText(f"[{timestamp}] {message}\n")


if __name__ == "__main__":
    app = wx.App(False)
    frame = ProtocolApp()
    app.MainLoop()